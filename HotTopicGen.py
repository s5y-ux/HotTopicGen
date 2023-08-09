import urllib.request
import json
import docx

currentHotTopic = str(input("Hot Topic Number>: "))
Name = str(input("Name [First, Last]>: "))
true_name = Name.split(",")
document_name = str(input("Name of document>: "))

#Function for calling API through URL
def Request(URLparameter: str) -> dict:

	#Stores JSON data in 'Response' variable
	Response = urllib.request.urlopen(URLparameter)

	#Parses data as dictionary and returns dict (Notice Return Data Type)
	Result = json.loads(Response.read())
	return(Result)

All_The_News = Request("https://newsdata.io/api/1/news?apikey=pub_261324ec320f89e25b0160ed20d8bad5267b4&q=politics")

Most_Recent_Article = All_The_News['results'][0]

document = docx.Document()

document.add_paragraph("{0} {1}".format(true_name[0], true_name[1]))
document.add_paragraph("Professor Callahan")
document.add_paragraph("American Government")
document.add_paragraph("Date")
document.add_paragraph("Mon/Wed 10:00am - 12:15pm")
document.add_heading("Hot Topic {0}".format(currentHotTopic))
document.add_paragraph(Most_Recent_Article["content"])
document.add_heading("Works Cited")
document.add_paragraph(Most_Recent_Article["link"])

document.save('{0}.docx'.format(document_name))

print(All_The_News)
